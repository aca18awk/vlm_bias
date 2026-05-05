import json
import os
import random

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

JSON_FILE    = "../15_Apr/vignettes_nature_paper/vignettes_nature_paper_20_Apr.json"
OUTPUT_DIR   = "../15_Apr/vignettes_nature_paper/docs_control"
MAPPING_FILE = "../15_Apr/vignettes_nature_paper/ground_truth_mapping_control.json"
IMAGES_DIR   = "../images"

# Conditions: (label, image_filename) — use None for no-photo (empty cell)
CONDITIONS = [
    ("grey_rect", "grey_rect.png"),
    ("no_photo", None),
]

# Existing CFD mapping to exclude codes from (avoids cross-file collisions)
EXISTING_MAPPING = "../15_Apr/vignettes_nature_paper/ground_truth_mapping.json"

os.makedirs(OUTPUT_DIR, exist_ok=True)

with open(JSON_FILE) as f:
    vignettes = json.load(f)

reserved_by_vignette: dict[str, set[str]] = {}
try:
    with open(EXISTING_MAPPING) as f:
        existing = json.load(f)
    for encoded in existing.values():
        prefix, code = encoded.split("_", 1)
        reserved_by_vignette.setdefault(prefix, set()).add(code)
except FileNotFoundError:
    pass

blind_mapping = {}

for i, vignette in enumerate(vignettes):
    vignette_num = i + 1
    prefix = f"{vignette_num:02d}"

    reserved = reserved_by_vignette.get(prefix, set())
    available = [str(c) for c in range(100, 1000) if str(c) not in reserved]
    codes = random.sample(available, len(CONDITIONS))
    code_iter = iter(codes)

    for label, img_file in CONDITIONS:
        doc = Document()

        title = doc.add_heading("EPIC - Patient Electronic Health Record", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        p = left_cell.paragraphs[0]
        p.add_run("PATIENT BANNER\n").bold = True
        p.add_run(f"Patient ID: {vignette['patient_id']}\n")
        p.add_run(f"Age: {vignette['age']}\n")
        p.add_run(f"Code Status: {vignette['code_status']}\n")
        p.add_run(f"Allergies: {vignette['allergies']}\n\n")

        p.add_run("CLINICAL VIGNETTE\n").bold = True
        p.add_run("Chief Complaint: ").bold = True
        p.add_run(f"{vignette['Chief Complaint']}\n")
        p.add_run("HPI: ").bold = True
        p.add_run(f"{vignette['HPI']}\n")
        p.add_run("Objective: ").bold = True
        p.add_run(f"{vignette['Objective']}\n")
        p.add_run("Diagnostics: ").bold = True
        p.add_run(f"{vignette['Diagnostics']}")

        p_img = right_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if img_file is not None:
            img_path = os.path.join(IMAGES_DIR, img_file)
            try:
                p_img.add_run().add_picture(img_path, width=Inches(1.5))
            except Exception:
                print(f"Warning: could not load image {img_path}")

        human_readable_name = f"{vignette_num:02d}_{label}"
        doc.save(os.path.join(OUTPUT_DIR, f"{human_readable_name}.docx"))

        encoded_name = f"{prefix}_{next(code_iter)}"
        blind_mapping[human_readable_name] = encoded_name

with open(MAPPING_FILE, "w") as f:
    json.dump(blind_mapping, f, indent=4)

print(f"Generated {len(blind_mapping)} documents in '{OUTPUT_DIR}'")
print(f"Mapping saved to '{MAPPING_FILE}'")
