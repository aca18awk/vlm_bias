import glob
import json
import os
import random
from collections import defaultdict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# 1. Setup paths
JSON_FILE = "../15_Apr/vignettes_nature_paper/vignettes_nature_paper_20_Apr.json"
OUTPUT_DIR = "../15_Apr/vignettes_nature_paper/docs_new_indian"
MAPPING_FILE = "../15_Apr/vignettes_nature_paper/ground_truth_mapping_new_indian.json"
CFD_DIR = "../images/CFD_indian"

os.makedirs(OUTPUT_DIR, exist_ok=True)

with open(JSON_FILE, "r") as f:
    vignettes = json.load(f)

# 2. Discover CFD images and group by subgroup code (AF, BM, WF, ...)
#    Filename format: CFD-{subgroup}-{id}-{age}-N.jpg
cfd_by_subgroup = defaultdict(list)
for path in sorted(glob.glob(os.path.join(CFD_DIR, "**", "CFD-*-N.jpg"), recursive=True)):
    fname = os.path.basename(path)
    parts = fname.split("-")
    if len(parts) >= 2:
        subgroup = parts[1]   # e.g. "AF", "WM"
        cfd_by_subgroup[subgroup].append(path)

print("CFD images found per subgroup:")
for sg, paths in sorted(cfd_by_subgroup.items()):
    print(f"  {sg}: {len(paths)}")

# 3. Build the flat list of (label, image_path_or_none) conditions per vignette
#    label format: {subgroup}_{face_idx}  e.g. "AF_1", "WM_3", "no_photo"
conditions = []
for subgroup, paths in sorted(cfd_by_subgroup.items()):
    for face_idx, img_path in enumerate(paths, start=1):
        conditions.append((f"{subgroup}_{face_idx}", img_path))
# conditions.append(("no_photo", None))

print(f"\nTotal conditions per vignette: {len(conditions)}")

# 4. Load existing codes to exclude (so new codes never collide with them)
EXISTING_MAPPING_FILE = "../15_Apr/vignettes_nature_paper/ground_truth_mapping.json"
try:
    with open(EXISTING_MAPPING_FILE) as f:
        existing = json.load(f)
    # existing codes keyed by vignette prefix, e.g. {"01": {"859"}, "02": {"769"}, ...}
    reserved_by_vignette: dict[str, set[str]] = {}
    for encoded in existing.values():
        prefix, code = encoded.split("_", 1)
        reserved_by_vignette.setdefault(prefix, set()).add(code)
except FileNotFoundError:
    reserved_by_vignette = {}

# Generate documents
blind_mapping = {}

for i, vignette in enumerate(vignettes):
    vignette_num = i + 1
    prefix = f"{vignette_num:02d}"

    # Draw unique codes for every condition in this vignette, excluding reserved ones
    reserved = reserved_by_vignette.get(prefix, set())
    available = [str(c) for c in range(100, 1000) if str(c) not in reserved]
    codes = random.sample(available, len(conditions))
    code_iter = iter(codes)

    for condition_label, img_path in conditions:
        doc = Document()

        title = doc.add_heading("EPIC - Patient Electronic Health Record", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)

        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # --- Left: clinical text ---
        p = left_cell.paragraphs[0]
        p.add_run("PATIENT BANNER\n").bold = True
        p.add_run(f"Patient ID: {vignette['patient_id']}\n")
        p.add_run(f"Age: {vignette['age']}\n")
        p.add_run(f"Code Status: {vignette['code_status']}\n")
        p.add_run(f"Allergies: {vignette['allergies']}\n\n")

        p.add_run("CLINICAL VIGNETTE\n").bold = True
        p.add_run("Chief Complaint: ").bold = True
        p.add_run(f"{vignette['Chief Complaint']}\n")
        p.add_run("HPI: ").bold = True
        p.add_run(f"{vignette['HPI']}\n")
        p.add_run("Objective: ").bold = True
        p.add_run(f"{vignette['Objective']}\n")
        p.add_run("Diagnostics: ").bold = True
        p.add_run(f"{vignette['Diagnostics']}")

        # --- Right: photo ---
        p_img = right_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if img_path is not None:
            try:
                p_img.add_run().add_picture(img_path, width=Inches(1.5))
            except Exception:
                print(f"Warning: could not load image {img_path}")

        human_readable_name = f"{vignette_num:02d}_{condition_label}"
        doc.save(os.path.join(OUTPUT_DIR, f"{human_readable_name}.docx"))

        encoded_name = f"{prefix}_{next(code_iter)}"
        blind_mapping[human_readable_name] = encoded_name

with open(MAPPING_FILE, "w") as f:
    json.dump(blind_mapping, f, indent=4)

print(f"\nGenerated {len(blind_mapping)} documents in '{OUTPUT_DIR}'")
print(f"Mapping saved to '{MAPPING_FILE}'")
