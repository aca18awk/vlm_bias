import json
import os
import random

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# 1. Setup paths and data
JSON_FILE = "../15_Apr/vignettes_nature_paper/vignettes_nature_paper_20_Apr.json"
OUTPUT_DIR = "../15_Apr/vignettes_nature_paper/docs_control"
MAPPING_FILE = "../15_Apr/vignettes_nature_paper/ground_truth_mapping_control.json"

# Added "no_photo" to the list of configurations
IMAGE_FILES = [
    # "black_man.png",
    # "black_woman.png",
    # "white_man.png",
    # "white_woman.png",
    "grey_rect.png"
    # "no_photo",
]

# Create output directory if it doesn't exist
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# Load the JSON data
with open(JSON_FILE, "r") as f:
    vignettes = json.load(f)

# Initialize the simplified mapping dictionary
blind_mapping = {}

# 2. Generate the documents
for i, vignette in enumerate(vignettes):
    vignette_num = i + 1

    for img_file in IMAGE_FILES:
        # Create a new Word Document
        doc = Document()

        # Add a title to make it look like an EHR System
        title = doc.add_heading("EPIC - Patient Electronic Health Record", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add a little spacing

        # Create a 1-row, 2-column table for our layout
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Set column widths (roughly 70% text, 30% image)
        table.columns[0].width = Inches(4.5)
        table.columns[1].width = Inches(2.0)

        # Get the left and right cells
        left_cell = table.cell(0, 0)
        right_cell = table.cell(0, 1)

        # --- LEFT CELL: TEXT ---
        # Patient Banner
        p = left_cell.paragraphs[0]
        p.add_run("PATIENT BANNER\n").bold = True
        p.add_run(f"Patient ID: {vignette['patient_id']}\n")
        p.add_run(f"Age: {vignette['age']}\n")
        p.add_run(f"Code Status: {vignette['code_status']}\n")
        p.add_run(f"Allergies: {vignette['allergies']}\n\n")

        # Clinical Vignette
        p.add_run("CLINICAL VIGNETTE\n").bold = True
        p.add_run("Chief Complaint: ").bold = True
        p.add_run(f"{vignette['Chief Complaint']}\n")

        p.add_run("HPI: ").bold = True
        p.add_run(f"{vignette['HPI']}\n")

        p.add_run("Objective: ").bold = True
        p.add_run(f"{vignette['Objective']}\n")

        p.add_run("Diagnostics: ").bold = True
        p.add_run(f"{vignette['Diagnostics']}")

        # --- RIGHT CELL: IMAGE ---
        p_img = right_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if img_file == "no_photo":
            demo_name = "no_photo"
            # We skip adding the picture, leaving the cell empty
        else:
            img_path = os.path.join("../images", img_file)
            demo_name = img_file.replace(".png", "")

            # Insert the image, locked to 1.5 inches wide
            try:
                p_img.add_run().add_picture(img_path, width=Inches(1.5))
            except Exception:
                print(
                    f"Warning: Could not load image {img_path}. Make sure it's in the same folder."
                )

        human_readable_name = f"{vignette_num:02d}_{demo_name}"
        filename = f"{human_readable_name}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)

        doc.save(filepath)

        # B. Generate the random encoded name for the future PDF/JPG
        # e.g., "01_492"
        random_code = str(random.randint(100, 999))
        encoded_base_name = f"{vignette_num:02d}_{random_code}"

        # Store the simplified key-value pair in our dictionary
        blind_mapping[human_readable_name] = encoded_base_name

# Write mapping to JSON
with open(MAPPING_FILE, "w") as f:
    json.dump(blind_mapping, f, indent=4)

print(f"Successfully generated Word documents in the '{OUTPUT_DIR}' folder!")
print(f"Mapping file saved as '{MAPPING_FILE}'!")
